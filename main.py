import os
from configparser import ConfigParser
from io import StringIO
from io import open
from concurrent.futures import ProcessPoolExecutor
from pathlib import Path

from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document


def read_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()

        content = return_str.getvalue()
        return_str.close()
        return content


def save_text_to_word(content, file_path):
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)


def pdf_to_word(pdf_file_path, word_file_path):
    content = read_from_pdf(pdf_file_path)
    save_text_to_word(content, word_file_path)


def main():
    config_parser = ConfigParser()
    config_parser.read('config.cfg')
    config = config_parser['default']

    tasks = []
    with ProcessPoolExecutor(max_workers=int(config['max_worker'])) as executor:
        input_folder = Path(config['input_folder'])
        for input_file in input_folder.iterdir():
            if input_file.suffix != '.pdf':
                print('File "{}" is not a valid pdf. Continuing...')
                continue

            output_file = Path(config['output_folder'], input_file.stem + '.doc')
            print('Processing file: "{}"'.format(input_file))
            result = executor.submit(pdf_to_word, input_file, output_file)
            tasks.append(result)

            while True:
                exit_flag = True
                for task in tasks:
                    if not task.done():
                        exit_flag = False
                if exit_flag:
                    print('Job Complete!')
                    raise SystemExit


if __name__ == '__main__':
    main()
